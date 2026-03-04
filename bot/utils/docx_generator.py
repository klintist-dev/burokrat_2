# bot/utils/docx_generator.py
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re
from datetime import datetime


class DocxGenerator:
    """Класс для создания Word-документов из текста"""

    @staticmethod
    def clean_text(text: str) -> str:
        """Очищает текст от Markdown-разметки"""
        # Убираем **жирный**
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        # Убираем *курсив*
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        # Убираем `код`
        text = re.sub(r'`(.*?)`', r'\1', text)
        # Убираем заголовки #
        text = re.sub(r'#{1,6}\s+', '', text)
        return text

    @staticmethod
    def create_document(title: str, content: str, user_id: int) -> str:
        """
        Создаёт Word-документ из текста
        Возвращает путь к файлу
        """
        # Создаём папку для документов, если её нет
        docs_dir = f"data/documents/{user_id}"
        os.makedirs(docs_dir, exist_ok=True)

        # Очищаем текст от Markdown
        clean_content = DocxGenerator.clean_text(content)

        # Создаём документ
        doc = Document()

        # Настройки стиля
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)

        # Заголовок документа
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(14)

        doc.add_paragraph()  # пустая строка

        # Добавляем дату создания
        date_str = datetime.now().strftime("%d.%m.%Y %H:%M")
        date_paragraph = doc.add_paragraph(f"Создано: {date_str}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_paragraph.runs[0].italic = True

        doc.add_paragraph()  # пустая строка

        # Добавляем основной текст
        # Разбиваем на параграфы по пустым строкам
        paragraphs = clean_content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())

        # Генерируем имя файла
        timestamp = int(datetime.now().timestamp())
        safe_title = re.sub(r'[^\w\s-]', '', title)[:30]
        filename = f"{safe_title}_{timestamp}.docx"
        filepath = os.path.join(docs_dir, filename)

        # Сохраняем
        doc.save(filepath)

        return filepath